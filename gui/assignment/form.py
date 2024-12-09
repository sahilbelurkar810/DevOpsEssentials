from tkinter import *
from docx import Document  # You'll need to install python-docx: pip install python-docx
from datetime import datetime

window = Tk()
window.geometry("1920x1080")
window.title("Registration Form")
window.configure(bg="white")
valx = 500
valy = 100
list1 = ["Name", "age", "Number", "EmailId", "College", "Course", "Usn", "Gender"]
list2 = list1

for i in list1:  
    lbl2 = Label(window, text=i, bg="white", fg="Black", font=("Roboto", 9))
    lbl2.place(x=valx, y=valy)
    valy += 50

valx = 600
valy = 100
entries = []  # Store Entry widgets in a list
for i in range(len(list1)):
    entry = Entry(window, width=30,bg="lightgrey")
    entry.place(x=valx, y=valy)
    entries.append(entry)  # Add Entry widget to list
    valy += 50

def send():
    try:
        # Try to open existing document
        try:
            doc = Document('form_data.docx')
            # Add a page break before new entries
            doc.add_page_break()
        except:
            # Create new document if it doesn't exist
            doc = Document()
        
        # Add a title with timestamp for this submission
        current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_heading(f'Form Submission - {current_time}', 0)
        
        # Add form data
        for label, entry in zip(list2, entries):
            value = entry.get()
            doc.add_paragraph(f"{label}: {value}")
            entry.delete(0, END)  # Clear the entry field
        
        # Save the document
        doc.save('form_data.docx')
        
        # Add a success message label
        success_label = Label(window, text="Data saved successfully!", fg="green")
        success_label.place(x=150, y=valy+100)
        # Remove the success message after 2 seconds
        window.after(2000, success_label.destroy)
        
    except Exception as e:
        # Add an error message label
        error_label = Label(window, text="Error saving data!", fg="red")
        error_label.place(x=150, y=valy+100)
        # Remove the error message after 2 seconds
        window.after(2000, error_label.destroy)

btn = Button(window, text="Send", command=send,width=30)
btn.place(x=530, y=valy+30)

window.mainloop()
